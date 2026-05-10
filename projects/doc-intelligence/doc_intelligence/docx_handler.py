"""Doc Intelligence - Word 文档处理"""

import os
from typing import Optional

try:
    from docx import Document
except ImportError:
    Document = None


class DocxHandler:
    """Word (.docx) 文档解析与处理"""

    def __init__(self):
        if Document is None:
            raise ImportError("请安装 python-docx: pip install python-docx")

    def extract_text(self, file_path: str) -> str:
        """提取 Word 文档全部文本"""
        doc = Document(file_path)
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text)
        return "\n".join(paragraphs)

    def extract_with_structure(self, file_path: str) -> str:
        """提取文本并保留标题层级"""
        doc = Document(file_path)
        parts = []
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            if p.style.name.startswith("Heading"):
                level = int(p.style.name.split()[-1]) if p.style.name.split()[-1].isdigit() else 1
                parts.append(f"{'#' * level} {p.text}")
            else:
                parts.append(p.text)
        return "\n\n".join(parts)

    def get_metadata(self, file_path: str) -> dict:
        """获取 Word 文档元信息"""
        doc = Document(file_path)
        core_properties = doc.core_properties
        return {
            "title": core_properties.title or "",
            "author": core_properties.author or "",
            "subject": core_properties.subject or "",
            "created": str(core_properties.created) if core_properties.created else "",
            "modified": str(core_properties.modified) if core_properties.modified else "",
            "paragraphs": len(doc.paragraphs),
        }

    def to_markdown(self, file_path: str, output_path: Optional[str] = None) -> str:
        """Word 转 Markdown"""
        md = self.extract_with_structure(file_path)
        if output_path:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(md)
        return md
